import chromadb
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

load_dotenv()

model = SentenceTransformer('all-MiniLM-L6-v2')
client = chromadb.Client()
collection = client.get_or_create_collection("resume")

def load_resume(file_path, file_type):
    text = ""
    if file_type == "pdf":
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
    elif file_type == "docx":
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    return text

def chunk_text(text, chunk_size=500):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append(chunk)
    return chunks

def index_resume(file_path, file_type):
    text = load_resume(file_path, file_type)
    chunks = chunk_text(text)
    embeddings = model.encode(chunks).tolist()
    collection.upsert(
        documents=chunks,
        embeddings=embeddings,
        ids=[f"chunk_{i}" for i in range(len(chunks))]
    )
    return f"Resume indexed: {len(chunks)} chunks stored"

def search_resume(query, n_results=3):
    query_embedding = model.encode([query]).tolist()
    results = collection.query(
        query_embeddings=query_embedding,
        n_results=n_results
    )
    return results["documents"][0]